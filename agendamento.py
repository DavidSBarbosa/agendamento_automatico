from tkinter.filedialog import askopenfilenames
from docx import Document
import pandas as pd
import ntpath
import re

diretorio_tabela = askopenfilenames()
for diretorio_atual in diretorio_tabela:
    nome_tabela = ntpath.basename(diretorio_atual)

    padrao_data = re.findall(r"([0-9]{8})\w", nome_tabela)
    padrao_data = ''.join(padrao_data)
    data_missa = padrao_data[:2] + '/' + padrao_data[2:4] + '/' + padrao_data[4:]

    padrao_hora = re.findall(r"([0-9][0-9][h]([0-9][0-9]|))", nome_tabela) #contem gambiarra
    hora_missa = ''.join(padrao_hora[0][0]) #contem gambiarra

    nome_do_arquivo = 'SANTA MISSA ' + padrao_data[:2] + '-' + padrao_data[2:4] + ' às ' + hora_missa

    df = pd.read_excel(diretorio_atual, usecols="C,D", skiprows=7)
    df = df["Nome"] + "" + df["Sobrenome"]

    df = df.dropna()

    df_list = df.values.tolist()

    document = Document()
    document.add_heading('Santa Missa ' + data_missa + " às " + hora_missa, 0)

    for x in df_list:
        document.add_paragraph(x, style ='List Bullet')

    document.save(nome_do_arquivo + '.docx')